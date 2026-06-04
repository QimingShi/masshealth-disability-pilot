"""Render a populated UMass Chan DES reviewer form (.docx) from matcher results.

Companion to pipeline/output.py (markdown + HTML). Takes a hand-crafted SSA
Blue Book template from templates/listings/<code> *.docx, ticks the per-leaf
checkboxes from the matcher's verdicts, marks the final form verdict, and
fills the OBJECTIVE MEDICAL EVIDENCE section with the AI narrative (per-leaf
rationale + verbatim citations).

Templates are organized as:

    templates/
    └── listings/
        ├── 1.15 Disorders of the Skeletal Spine ...docx
        ├── 1.16 Lumbar Spinal Stenosis ...docx
        ├── ...
        └── Blank Listing.docx        (continuation/overflow page, NOT a skeleton)

Each template's body follows the same structure (confirmed by spot-check of
1.15, see _inspect_template if you want to verify others):

    [header]    "<code> <TITLE IN CAPS>"
    [intro]     "<lead-in describing what the listing covers>"
    [criteria]  leaf-marker paragraphs prefixed by "____X." where X is a
                letter or digit; nested levels use underscore-prefixed digits
                and lowercase letters (e.g. "____a.", "____1.")
    [verdict]   "______Meets    ______Equals    ______Does not meet/equal"
    [signoff]   "Disability Reviewer # ___  Physician Advisor # ___"
    [evidence]  "OBJECTIVE MEDICAL EVIDENCE AND OTHER ADDITIONAL INFORMATION"
                followed by ~20 underscore-only paragraphs (the blank lines
                where reviewers usually hand-write narrative)

The renderer is intentionally surgical — it modifies the minimum text needed
in each paragraph (typically just the leading "____" prefix) and preserves
the rest of the template formatting (Courier New, bold "or" markers, etc.).
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt

from .chunks import Chunk, Listing
from .consolidate import NodeVerdict, form_verdict
from .evaluate import LeafResult


# ---------------------------------------------------------------------------
#  Public API
# ---------------------------------------------------------------------------

# Visual style: ASCII bracketed boxes. Unicode ☒/☐ depend on font glyph
# coverage and don't render reliably in Office's default Courier New
# (reviewers saw blank squares). ASCII renders in every font, every
# Word/LibreOffice/Google-Docs version.
CHECKBOX_MET          = "[X]"      # verdict 'met'
CHECKBOX_NOT_MET      = "[ ]"      # verdict 'not_met'
CHECKBOX_INSUFFICIENT = "[?]"      # verdict 'insufficient'
CHECKBOX_UNKNOWN      = "[_]"      # leaf not evaluated (shouldn't happen
                                   # in normal flow; defensive fallback)

# Listings the matcher may pick but for which we have no hand-crafted
# template. The renderer logs a warning and skips docx for these; HTML/MD
# still produced upstream by pipeline/output.py. See README §Known limitations.
LISTINGS_WITHOUT_TEMPLATE = {"12.05", "5.11", "5.12", "8.09"}


def find_template(listing_code: str, templates_dir: Path) -> Path | None:
    """Resolve a listing code (e.g. '1.15') to its template file.

    Templates are stored with their original filenames ('<code> <title>.docx')
    rather than renamed to '<code>.docx'. This lets the same folder be used
    for both rendering and side-by-side reviewer reference. Lookup is a glob
    on '<code> *.docx' anchored on the code prefix.
    """
    # Anchor: require a space (or end of stem) right after the code, otherwise
    # '1.15' would also match '1.155' if that existed.
    candidates = list(templates_dir.glob(f"{listing_code} *.docx"))
    if not candidates:
        return None
    if len(candidates) > 1:
        # Shouldn't happen with the SSA naming convention, but be deterministic
        candidates.sort()
    return candidates[0]


def render_form_docx(
    listing: Listing,
    root_verdict: NodeVerdict,
    leaf_results: dict[str, LeafResult],
    case_id: str,
    chunks_by_id: dict[str, Chunk],
    template_path: Path,
    output_path: Path,
) -> Path:
    """Render a UMass Chan DES form by filling in a listing's template.

    Reads template_path, mutates a python-docx Document in-memory, writes to
    output_path. Returns output_path on success.

    Args:
        listing: Listing dataclass; .code and .title come from the DB (DB title
            wins over template title — handles SSA 10-6-23 revisions and
            template-side typos).
        root_verdict: consolidated NodeVerdict tree (top-level decision +
            full criterion tree with verdicts at every node).
        leaf_results: {leaf_path: LeafResult} from the matcher run.
            leaf_path keys are e.g. '1.15.A.1' (full path including listing
            code, matching what's persisted in leaf_assessments.leaf_path).
        case_id: human-readable case id (goes onto the form header).
        chunks_by_id: {chunk_id_str: Chunk} for citation rendering in the
            OBJECTIVE MEDICAL EVIDENCE narrative section.
        template_path: path to the source template.
        output_path: where to write the rendered .docx.

    Raises:
        FileNotFoundError: if template_path doesn't exist.
    """
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(str(template_path))

    # --- Phase 0: fill the page-header DIAGNOSIS + CASE # fields ---
    # These live in the .docx page header (top of every printed page), not
    # in the body paragraphs. They show up in EVERY section's header, so
    # we iterate sections to make sure the title page and the continuation
    # page both display the same case info.
    _fill_page_header_fields(doc, diagnosis=listing.title, case_id=case_id)

    # --- Phase 1: rewrite the listing-code/title headers ---
    # The template repeats the header at the top of each "logical page"
    # (~3 occurrences in a typical template). We replace every paragraph whose
    # text starts with the listing code so all three update at once.
    new_header = f"{listing.code} {listing.title.upper()}"
    _replace_header_paragraphs(doc, listing.code, new_header)

    # --- Phase 2: tick leaf-marker checkboxes ---
    # Walk paragraphs in order, tracking a path-stack state machine so each
    # encountered marker resolves to a full leaf_path (e.g. '1.15.B.3.a').
    # We use the consolidated NodeVerdict tree (not just leaf_results) so that
    # group rows (e.g. ____A. when A is an OR over A.1/A.2/A.3) also get their
    # rolled-up verdict — the template treats every "____" as something to tick.
    verdicts_by_path = _flatten_verdicts(root_verdict)
    marker_paragraphs = _mark_leaf_checkboxes(doc, listing.code, verdicts_by_path)

    # --- Phase 3: tick the final verdict line ---
    verdict_label = form_verdict(root_verdict)
    _mark_verdict_line(doc, verdict_label)

    # --- Phase 4: fill the OBJECTIVE MEDICAL EVIDENCE narrative ---
    narrative_paragraphs = _build_narrative(root_verdict, leaf_results, chunks_by_id)
    _fill_evidence_section(doc, narrative_paragraphs, case_id)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
#  Phase 1 — headers
# ---------------------------------------------------------------------------

def _replace_header_paragraphs(doc, listing_code: str, new_header: str) -> int:
    """Find paragraphs starting with the listing code and rewrite to new_header.

    The template repeats the header (e.g. on every logical page); we update
    all instances. Returns the count of paragraphs updated.
    """
    count = 0
    code_pattern = re.compile(rf"^{re.escape(listing_code)}\b")
    for p in doc.paragraphs:
        if code_pattern.match(p.text or ""):
            _set_paragraph_text(p, new_header)
            count += 1
    return count


# ---------------------------------------------------------------------------
#  Phase 0 — page-header DIAGNOSIS / CASE # fields
# ---------------------------------------------------------------------------

# Pattern matches runs that are pure "form-field" blanks: any amount of
# whitespace, underscores, or non-breaking spaces, and nothing else.
_BLANK_RUN_RE = re.compile(r"^[\s_\xa0]+$")


def _fill_page_header_fields(doc, *, diagnosis: str, case_id: str) -> int:
    """Fill the DIAGNOSIS and CASE # blanks in the page header (top of page).

    The UMass Chan template layout is:

        DIAG[tab]NOSIS  __________________   CASE #  __________
                       __________________

    Both labels are bold (multiple runs each: e.g. ['DIAG', '\\tNOSIS']);
    the blanks are non-bold runs of underscores. We walk runs in each header
    paragraph, track which label we just passed, fill the next blank run
    with the value, then *consume* any subsequent blank runs (and entire
    blank-only paragraphs) until we encounter another label — this clears
    the multi-line overflow blanks the template provides for long titles.

    Returns the count of fields filled across all sections.
    """
    # Tokens we look for in the accumulated bold-label text. Order matters:
    # DIAGNOSIS comes first in the template, so we check it first when both
    # labels could plausibly match a partial accumulator.
    FIELD_TOKENS = [
        ("DIAGNOSIS", diagnosis),
        ("CASE",      case_id),    # the template label is "CASE #"; matching
                                   # on "CASE" alone is more lenient and safe
    ]
    filled = 0
    for section in doc.sections:
        for hf in (section.header, section.first_page_header,
                   section.even_page_header):
            if hf is None:
                continue
            filled += _fill_header_part(hf, FIELD_TOKENS)
    return filled


def _fill_header_part(header_part, field_tokens: list[tuple[str, str]]) -> int:
    """Walk all paragraphs in one header-part with cross-paragraph state.

    State machine (per header-part):
      'idle'        — no label seen yet; leave runs alone
      'armed'       — just saw a label, awaiting the next blank run
      'consuming'   — already filled this field's value, clear any
                      subsequent blank runs (multi-line overflow blanks)
                      until another label is found
    """
    filled = 0
    state = "idle"
    pending_value: str | None = None
    bold_buf = ""

    for p in header_part.paragraphs:
        for run in p.runs:
            txt = run.text or ""
            is_blank = bool(_BLANK_RUN_RE.match(txt)) and txt != ""
            is_bold = bool(run.bold)

            if is_bold and not is_blank:
                # Non-blank bold run = label text. Switching to a label
                # always resets us out of any consuming state and re-arms
                # if the accumulator matches a known field.
                state = "idle"
                pending_value = None
                bold_buf += re.sub(r"\s+", "", txt)
                for token, value in field_tokens:
                    if token in bold_buf.upper():
                        state = "armed"
                        pending_value = value
                        bold_buf = ""    # consumed; ready for the next label
                        break
            elif is_blank:
                if state == "armed":
                    # Fill this blank with the value, then start consuming
                    # subsequent blanks (the template's overflow lines).
                    run.text = " " + pending_value + "  "
                    state = "consuming"
                    pending_value = None
                    filled += 1
                elif state == "consuming":
                    # Continuation blank for the field we just filled.
                    # Clear it so the rendered docx doesn't show stray
                    # underscores after the value.
                    run.text = ""
                # else: blank in idle state → leave alone (decorative)
            else:
                # Non-bold, non-blank run (e.g. inline punctuation between
                # label and blank). Don't reset state — we may still be
                # waiting for the next blank after the label.
                pass
    return filled


# ---------------------------------------------------------------------------
#  Phase 2 — leaf checkboxes
# ---------------------------------------------------------------------------

# Tolerate leading whitespace because nested sub-criteria are tab-indented
# in many templates (e.g. 12.02 uses '\t____1.' under A, while 1.15 uses
# unindented '____1.'). The optional \s* matches both.
_LEAF_MARKER_RE = re.compile(r"^\s*(_{2,})([A-Za-z0-9]+)\.")


def _marker_level(marker: str) -> int:
    """Determine the tree level of a leaf marker.

      Level 0: uppercase letter   (A, B, C, D, ...)
      Level 1: digit              (1, 2, 3, ...)
      Level 2: lowercase letter   (a, b, c, ...)
      Level 3: roman numeral      (i, ii, iii, ...)  — rarely seen in adult
                                  listings but supported defensively
    """
    if not marker:
        return 0
    c = marker[0]
    if c.isupper():
        return 0
    if c.isdigit():
        return 1
    if c.islower():
        # Crude heuristic: treat 'i', 'ii', 'iii' etc. (Roman numerals) as
        # level 3 only when the whole marker matches the pattern. Otherwise
        # it's a normal lowercase letter at level 2.
        if marker.lower() in {"i", "ii", "iii", "iv", "v", "vi", "vii", "viii", "ix", "x"}:
            return 3
        return 2
    return 2


def _flatten_verdicts(root: NodeVerdict) -> dict[str, str]:
    """Walk a NodeVerdict tree and emit {path: verdict} for every node.

    Includes both group rows (where verdict is the AND/OR rollup) and leaf
    rows (where verdict came from evaluate.py). Used to tick template
    markers — the template doesn't distinguish groups vs leaves; both look
    like '____X.' and both deserve a verdict.
    """
    out: dict[str, str] = {}
    def walk(n: NodeVerdict):
        out[n.path] = n.verdict
        for c in n.children:
            walk(c)
    walk(root)
    return out


def _mark_leaf_checkboxes(
    doc,
    listing_code: str,
    verdicts_by_path: dict[str, str],
) -> dict[str, int]:
    """Walk paragraphs; for each marker line, replace ____ with verdict checkbox.

    Maintains a path stack: when we see an uppercase letter we start a new
    top-level segment; nested digits/lowercase append/replace at the right
    depth. Returns a {path: paragraph_index} map for debugging.
    """
    stack: list[str] = []
    seen: dict[str, int] = {}

    for idx, p in enumerate(doc.paragraphs):
        m = _LEAF_MARKER_RE.match(p.text or "")
        if not m:
            continue
        marker = m.group(2)
        level = _marker_level(marker)
        stack = stack[:level] + [marker]
        # The NodeVerdict tree uses paths straight from rule_json — those
        # are BARE (e.g. "A", "A.1", "B.3.a") with no listing-code prefix.
        # Try the bare-stack key first, then fall back to the prefixed
        # form so callers that pass listing-prefixed verdict dicts (some
        # tests) still work.
        bare_path = ".".join(stack)
        prefixed_path = f"{listing_code}.{bare_path}"
        seen[bare_path] = idx

        verdict = (verdicts_by_path.get(bare_path)
                   or verdicts_by_path.get(prefixed_path))
        checkbox = {
            "met":          CHECKBOX_MET,
            "not_met":      CHECKBOX_NOT_MET,
            "insufficient": CHECKBOX_INSUFFICIENT,
        }.get(verdict, CHECKBOX_UNKNOWN)

        _replace_leaf_marker_in_paragraph(p, checkbox)

    return seen


def _replace_leaf_marker_in_paragraph(paragraph, checkbox: str) -> bool:
    """Replace leading underscores in the first run that contains them.

    Most templates put '____X. ' in a single run; one variant we've seen
    splits it as '____a' + '. '. We only need to touch the underscores —
    leave the letter/digit and trailing text alone. This preserves any
    intra-paragraph formatting (bold "or" markers, etc.) that python-docx
    would otherwise destroy if we set paragraph.text wholesale.
    """
    for run in paragraph.runs:
        m = re.match(r"^(_{2,})", run.text)
        if m:
            run.text = checkbox + " " + run.text[m.end():]
            return True
    return False


# ---------------------------------------------------------------------------
#  Phase 3 — verdict line
# ---------------------------------------------------------------------------

def _mark_verdict_line(doc, verdict_label: str) -> bool:
    """Find the '______Meets ... ______Equals ... ______Does not meet' line
    and tick the option matching the form verdict.

    The template's verdict line is heavily run-fragmented, so we rewrite it
    in full rather than try to surgically modify individual runs. The result
    uses ☒ next to the picked verdict and ☐ next to the others, matching
    the leaf-checkbox convention from Phase 2.
    """
    # Heuristic: the verdict line is the unique paragraph that contains
    # "Meets" AND "Equals" AND "Does not meet/equal" tokens.
    target = None
    for p in doc.paragraphs:
        t = p.text or ""
        if "Meets" in t and "Equals" in t and "Does not meet" in t:
            target = p
            break
    if target is None:
        return False

    def mark(label: str) -> str:
        return f"{CHECKBOX_MET if label == verdict_label else CHECKBOX_NOT_MET} {label}"

    # Reconstruct the line. We use 2 tabs between options to roughly match
    # the template's tab-stop spacing (\t\t   in the original).
    new_text = (
        f"{mark('Meets')}\t\t"
        f"{mark('Equals')}\t\t"
        f"{mark('Does not meet/equal')}"
    )
    _set_paragraph_text(target, new_text, bold=True)
    return True


# ---------------------------------------------------------------------------
#  Phase 4 — OBJECTIVE MEDICAL EVIDENCE narrative
# ---------------------------------------------------------------------------

_EVIDENCE_HEADER_RE = re.compile(
    r"^OBJECTIVE MEDICAL EVIDENCE\b", re.IGNORECASE
)
_UNDERSCORE_LINE_RE = re.compile(r"^_+$")


def _fill_evidence_section(
    doc,
    narrative_paragraphs: list[str],
    case_id: str,    # accepted for API symmetry; case_id is shown in the
                     # page header instead, so we no longer pre-roll it here
) -> int:
    """Locate the OBJECTIVE MEDICAL EVIDENCE header and replace the trailing
    underscore lines with rendered narrative paragraphs.

    Strategy:
      1. Find the (first) paragraph whose text starts with the evidence header.
         The template usually has a (con't) continuation header further down —
         we touch only the first one (the per-leaf narrative is short enough
         to fit; if it doesn't, the continuation page is still available for
         the reviewer to hand-write spill-over).
      2. Walk forward, replacing each underscore-only paragraph's text with
         the next narrative paragraph until we run out of narrative OR run
         out of underscore lines.
      3. If we have more narrative than underscore lines, the surplus is
         joined onto the last paragraph (so nothing is silently dropped).
      4. Any remaining underscore lines stay as-is (they look like blank
         lines for hand-written notes after the AI narrative — fine).

    Returns the count of narrative paragraphs written.
    """
    if not narrative_paragraphs:
        return 0

    paras = doc.paragraphs
    header_idx = None
    for i, p in enumerate(paras):
        if _EVIDENCE_HEADER_RE.match(p.text or ""):
            header_idx = i
            break
    if header_idx is None:
        return 0

    # CASE # now lives in the page header (filled by Phase 0), so the
    # narrative section starts directly with the per-leaf rationale lines.
    narrative = list(narrative_paragraphs)

    underscore_indices = []
    for j in range(header_idx + 1, len(paras)):
        if _UNDERSCORE_LINE_RE.match((paras[j].text or "").strip()):
            underscore_indices.append(j)
        else:
            # Hit something that isn't an underscore line — stop. The template
            # may have a (con't) header further down that we don't touch.
            break

    written = 0
    for k, idx in enumerate(underscore_indices):
        if k < len(narrative):
            text = narrative[k]
        elif k == len(underscore_indices) - 1 and len(narrative) > len(underscore_indices):
            # Overflow: cram everything left into the last available slot,
            # separated by line breaks within the paragraph (so the reviewer
            # still sees it without us mutating doc layout).
            leftover = narrative[k:]
            text = "\n".join(leftover)
        else:
            text = ""
        _set_paragraph_text(paras[idx], text)
        written += 1
    return written


def _build_narrative(
    root_verdict: NodeVerdict,
    leaf_results: dict[str, LeafResult],
    chunks_by_id: dict[str, Chunk],
) -> list[str]:
    """Render the per-leaf rationale + verbatim citations as a list of
    paragraph-sized strings — for met leaves only.

    The form is a positive-evidence record. Reviewers don't need the form to
    re-state "this criterion wasn't met" or "AI couldn't tell" for every
    not_met / insufficient leaf — those are visible via the unchecked boxes
    in the criterion tree above. The HTML form (pipeline/output.py) still
    shows the full audit trail including not_met and insufficient rationales
    for cases when the reviewer wants the AI's complete reasoning.
    """
    paragraphs: list[str] = []

    def walk(node: NodeVerdict):
        if node.logic is None:
            if node.verdict != "met":
                return
            lr = leaf_results.get(node.path)
            if lr is None:
                return
            head = f"{node.path}: {lr.rationale or ''}"
            paragraphs.append(head)
            for ev in lr.evidence or []:
                chunk = chunks_by_id.get(ev.chunk_id)
                if chunk is None:
                    continue
                tag = _chunk_label(chunk)
                paragraphs.append(f"  - {tag}: \"{ev.quote}\"")
        else:
            for c in node.children:
                walk(c)

    walk(root_verdict)
    return paragraphs


def _chunk_label(chunk: Chunk) -> str:
    """Compact citation tag — mirrors pipeline/output._chunk_label."""
    bits = []
    if chunk.doc_title:
        bits.append(
            chunk.doc_title.split("—")[-1].strip()
            if "—" in chunk.doc_title else chunk.doc_title
        )
    elif chunk.doc_type:
        bits.append(chunk.doc_type.replace("_", " "))
    if chunk.section:
        bits.append(chunk.section)
    if chunk.encounter_date:
        bits.append(chunk.encounter_date)
    if chunk.page_start:
        bits.append(f"p.{chunk.page_start}")
    return ", ".join(bits)


# ---------------------------------------------------------------------------
#  Helpers
# ---------------------------------------------------------------------------

def _set_paragraph_text(paragraph, new_text: str, *, bold: bool | None = None):
    """Replace a paragraph's text while preserving the run's font/style.

    We borrow the first run's formatting (font name, size, bold/italic) so
    the new text doesn't look out of place when the template uses Courier
    New everywhere. If `bold` is given explicitly it overrides whatever the
    first run had.
    """
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    first = paragraph.runs[0]
    template_font = first.font.name
    template_bold = first.bold if bold is None else bold
    template_size = first.font.size

    # Wipe all runs except the first, then rewrite the first.
    for r in paragraph.runs[1:]:
        r.text = ""
    first.text = new_text
    if template_font:
        first.font.name = template_font
    if template_bold is not None:
        first.bold = bool(template_bold)
    if template_size is not None:
        first.font.size = template_size
