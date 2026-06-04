"""Rewrite already-rendered docx files in output/ to the v2 conventions:
ASCII checkboxes + met-only narrative, without re-running the matcher.

Use this when you've changed the docx renderer and want existing case
bundles to pick up the new format for $0 (instead of paying $1-3 per case
to re-run `py run.py --from-db`).

Run from the repo root:

    py backfill_docx.py

What it does to each output/<case_id>/*.docx:
  1. Replace Unicode boxes in leaf-marker paragraphs:
       ☒ A. -> [X] A.    (met)
       ☐ 1. -> [ ] 1.    (not_met)
       ? 4.  -> [?] 4.    (insufficient)
  2. Replace the verdict line (Meets/Equals/Does not meet/equal):
       ☒ Meets ... ☐ Equals ... -> [X] Meets ... [ ] Equals ...
  3. In the OBJECTIVE MEDICAL EVIDENCE narrative:
       - keep met-leaf lines, rewrite "1.15.A.1 (met): X" -> "1.15.A.1: X"
       - blank out not_met and insufficient header lines AND their
         indented continuation citation lines
       - replace Unicode bullet '•' and smart quotes with ASCII equivalents

Idempotent: re-running on already-converted files makes zero changes.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document


HERE = Path(__file__).parent
OUTPUT_ROOT = HERE / "output"

LEAF_MARKER_OLD_RE = re.compile(r"^([☒☐?])\s+([A-Za-z0-9]+)\.")
#                                  ☒        ☐
VERDICT_LINE_RE = re.compile(r"Meets.*Equals.*Does not meet", re.DOTALL)
NARRATIVE_HEADER_RE = re.compile(
    r"^(\d+\.\d+(?:\.[A-Za-z0-9]+)+)\s+\((met|not_met|insufficient)\):\s*(.*)$"
)
CONTINUATION_RE = re.compile(r"^\s+[•\-]")   # • or - bullet

BOX_MAP = {"☒": "[X]", "☐": "[ ]", "?": "[?]"}


def _set_paragraph_text(paragraph, new_text: str):
    """Mirror of pipeline.output_docx._set_paragraph_text — preserve the
    first run's font, clear other runs, set the text."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    first = paragraph.runs[0]
    for r in paragraph.runs[1:]:
        r.text = ""
    first.text = new_text


def backfill_one(path: Path) -> int:
    doc = Document(str(path))
    changes = 0
    skip_continuation = False

    for p in doc.paragraphs:
        text = p.text or ""

        if not text.strip():
            skip_continuation = False
            continue

        # 1. Leaf-marker paragraph (body criterion tree)
        m = LEAF_MARKER_OLD_RE.match(text)
        if m:
            old_box = m.group(1)
            new_box = BOX_MAP.get(old_box)
            if new_box:
                # Replace the old box character in the first run that has it.
                for run in p.runs:
                    if old_box in run.text:
                        run.text = run.text.replace(old_box, new_box, 1)
                        changes += 1
                        break
            skip_continuation = False
            continue

        # 2. Verdict line — heavily fragmented, easier to rewrite whole text
        if VERDICT_LINE_RE.search(text) and ("☒" in text or "☐" in text):
            new = text.replace("☒", "[X]").replace("☐", "[ ]")
            _set_paragraph_text(p, new)
            changes += 1
            skip_continuation = False
            continue

        # 3. Narrative header — '<path> (<verdict>): <rationale>'
        m = NARRATIVE_HEADER_RE.match(text)
        if m:
            leaf_path, verdict, rationale = m.group(1), m.group(2), m.group(3)
            if verdict == "met":
                # New format drops the "(met)" tag — boxes already say "met".
                new = f"{leaf_path}: {rationale}"
                if new != text:
                    _set_paragraph_text(p, new)
                    changes += 1
                skip_continuation = False
            else:
                # not_met / insufficient — drop this block.
                _set_paragraph_text(p, "")
                changes += 1
                skip_continuation = True
            continue

        # 4. Continuation citation line (indented bullet)
        if CONTINUATION_RE.match(text):
            if skip_continuation:
                _set_paragraph_text(p, "")
                changes += 1
            else:
                # Cosmetic: replace Unicode bullet/quotes with ASCII.
                new = (text
                       .replace("•", "-")
                       .replace("“", '"')
                       .replace("”", '"'))
                if new != text:
                    _set_paragraph_text(p, new)
                    changes += 1
            continue

        # Other content (headers, "AND", static template text) — leave alone
        skip_continuation = False

    if changes:
        doc.save(str(path))
    return changes


def main():
    if not OUTPUT_ROOT.exists():
        print(f"No output/ directory at {OUTPUT_ROOT}")
        return 0

    docx_files = list(OUTPUT_ROOT.rglob("*.docx"))
    if not docx_files:
        print(f"No .docx files under {OUTPUT_ROOT}")
        return 0

    total_changes = 0
    print(f"Backfilling {len(docx_files)} docx file(s):")
    for p in docx_files:
        try:
            n = backfill_one(p)
        except Exception as e:
            print(f"  ERROR {p.relative_to(HERE)}: {e}")
            continue
        rel = p.relative_to(HERE)
        if n:
            print(f"  {rel}: {n} change(s)")
        else:
            print(f"  {rel}: (no changes needed)")
        total_changes += n
    print(f"\nTotal changes across all files: {total_changes}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
